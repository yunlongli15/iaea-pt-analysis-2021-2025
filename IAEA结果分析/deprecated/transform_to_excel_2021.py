from docx import Document
import pandas as pd
import os
import re

def extract_and_merge_labcode_tables(docx_path, output_folder="merged_labcode_tables_2021"):
    """
    提取并合并包含Labcode的表格
    
    规则：
    1. 只提取表头包含'labcode'的表格
    2. 连续表格中，如果下一个表格的首行labcode > 上一个表格的末行labcode，则合并
    3. 否则开始新表格
    """
    
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(docx_path)
    
    print(f"开始处理文档: {os.path.basename(docx_path)}")
    print(f"文档中总共有 {len(doc.tables)} 个表格对象")
    
    # 存储所有找到的表格数据
    all_valid_tables = []
    
    # 第一步：提取所有包含labcode的表格
    for table_idx, table in enumerate(doc.tables):
        # 提取表格数据
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(cell_text for cell_text in row_data):
                table_data.append(row_data)
        
        if not table_data:
            continue
        
        # 检查表头是否包含'labcode'（不区分大小写）
        header_row = ' '.join(table_data[0]).lower()
        if 'labcode' in header_row:
            print(f"表格 {table_idx+1}: 找到包含labcode的表格 ({len(table_data)}行)")
            all_valid_tables.append({
                'index': table_idx,
                'data': table_data,
                'first_labcode': extract_labcode_from_row(table_data[1] if len(table_data) > 1 else []),
                'last_labcode': extract_labcode_from_row(table_data[-1] if len(table_data) > 1 else [])
            })
        else:
            print(f"表格 {table_idx+1}: 跳过（不包含labcode）")
    
    print(f"\n找到 {len(all_valid_tables)} 个包含labcode的表格")
    
    # 第二步：合并连续的表格
    merged_tables = []
    current_table = None
    
    for table_info in all_valid_tables:
        data = table_info['data']
        first_labcode = table_info['first_labcode']
        
        if current_table is None:
            # 开始第一个表格
            current_table = {
                'data': data,
                'first_labcode': first_labcode,
                'last_labcode': table_info['last_labcode']
            }
        else:
            # 检查是否可以合并
            can_merge = False
            
            # 情况1：当前表格的第一行是表头，但第二行开始是数据
            if len(data) > 1:
                data_first_labcode = extract_labcode_from_row(data[1])
                
                # 如果当前表格的数据行labcode大于上一个表格的最后labcode，则合并
                if (data_first_labcode is not None and 
                    current_table['last_labcode'] is not None and
                    is_numeric_labcode(data_first_labcode) and 
                    is_numeric_labcode(current_table['last_labcode'])):
                    
                    try:
                        data_first_num = int(data_first_labcode)
                        current_last_num = int(current_table['last_labcode'])
                        if data_first_num > current_last_num:
                            can_merge = True
                            print(f"  可以合并：当前labcode={data_first_num} > 上一个labcode={current_last_num}")
                    except:
                        pass
            
            # 情况2：当前表格没有表头，直接是数据行
            elif len(data) == 1 or not any('labcode' in cell.lower() for cell in data[0]):
                first_data_labcode = extract_labcode_from_row(data[0])
                if (first_data_labcode is not None and 
                    current_table['last_labcode'] is not None and
                    is_numeric_labcode(first_data_labcode) and 
                    is_numeric_labcode(current_table['last_labcode'])):
                    
                    try:
                        data_first_num = int(first_data_labcode)
                        current_last_num = int(current_table['last_labcode'])
                        if data_first_num > current_last_num:
                            can_merge = True
                            print(f"  可以合并：当前labcode={data_first_num} > 上一个labcode={current_last_num}")
                    except:
                        pass
            
            if can_merge:
                # 合并表格（跳过重复的表头）
                # 检查当前表格的第一行是否是表头
                first_row_is_header = any('labcode' in cell.lower() for cell in data[0])
                
                if first_row_is_header and len(data) > 1:
                    # 跳过表头，只合并数据行
                    current_table['data'].extend(data[1:])
                else:
                    # 直接合并所有行
                    current_table['data'].extend(data)
                
                # 更新最后一个labcode
                if data:
                    current_table['last_labcode'] = extract_labcode_from_row(data[-1])
            else:
                # 保存当前表格，开始新表格
                merged_tables.append(current_table)
                current_table = {
                    'data': data,
                    'first_labcode': first_labcode,
                    'last_labcode': table_info['last_labcode']
                }
    
    # 添加最后一个表格
    if current_table is not None:
        merged_tables.append(current_table)
    
    print(f"\n合并后得到 {len(merged_tables)} 个表格")
    
    # 第三步：保存表格
    saved_count = 0
    for i, table_info in enumerate(merged_tables, 1):
        table_data = table_info['data']
        
        if not table_data:
            continue
        
        # 创建DataFrame
        df = pd.DataFrame(table_data)
        
        # 生成文件名
        # 尝试从表格中提取描述信息
        table_name = f"Labcode_Table_{i}"
        
        # 检查表格第一行（通常是表头）寻找描述信息
        if len(table_data) > 0:
            header_text = ' '.join(table_data[0])
            if 'Eu-152' in header_text:
                table_name = f"Eu152_Labcode_{i}"
            elif 'Po-210' in header_text:
                table_name = f"Po210_Labcode_{i}"
            elif 'Pb-210' in header_text:
                table_name = f"Pb210_Labcode_{i}"
            elif 'Total U' in header_text or 'Total uranium' in header_text.lower():
                table_name = f"TotalU_Labcode_{i}"
        
        # 添加统计信息
        first_labcode = table_info['first_labcode'] or "?"
        last_labcode = table_info['last_labcode'] or "?"
        
        # 安全文件名
        safe_name = re.sub(r'[\\/*?:"<>|]', "_", table_name)
        excel_path = os.path.join(output_folder, f"{i:03d}_{safe_name}_{first_labcode}-{last_labcode}.xlsx")
        
        # 保存Excel
        df.to_excel(excel_path, index=False, header=False)
        saved_count += 1
        
        print(f"保存表格 {i}: {safe_name} ({len(table_data)}行, Labcode范围: {first_labcode}-{last_labcode})")
    
    print(f"\n成功保存 {saved_count} 个表格到: {output_folder}")
    return saved_count


def extract_labcode_from_row(row):
    """从行数据中提取labcode（假设labcode在第一列）"""
    if not row or len(row) == 0:
        return None
    
    # labcode通常在第一列
    labcode_cell = row[0].strip()
    
    # 尝试提取数字labcode
    # 匹配数字，可能是整数或小数
    match = re.search(r'^\s*(\d+(?:\.\d+)?)\s*$', labcode_cell)
    if match:
        return match.group(1)
    
    # 如果直接就是数字，返回
    if labcode_cell.isdigit():
        return labcode_cell
    
    # 尝试其他格式
    if labcode_cell:
        # 移除可能的标点符号
        clean_code = re.sub(r'[^\d\.]', '', labcode_cell)
        if clean_code:
            return clean_code
    
    return None


def is_numeric_labcode(labcode):
    """检查labcode是否是数字"""
    if labcode is None:
        return False
    
    # 尝试转换为整数或浮点数
    try:
        float(labcode)
        return True
    except:
        return False


# 更简单的实现：直接按labcode顺序合并
def simple_merge_by_labcode(docx_path, output_folder="simple_merged_tables"):
    """
    简单方法：按labcode顺序合并表格
    """
    
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(docx_path)
    
    print(f"处理文档: {os.path.basename(docx_path)}")
    
    # 第一步：收集所有包含labcode的表格数据
    all_table_data = []
    
    for table_idx, table in enumerate(doc.tables):
        table_rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(cell_text for cell_text in row_data):
                table_rows.append(row_data)
        
        if not table_rows:
            continue
        
        # 检查是否包含labcode
        first_row_text = ' '.join(table_rows[0]).lower()
        if 'labcode' in first_row_text:
            all_table_data.extend(table_rows)
            print(f"表格 {table_idx+1}: 添加 {len(table_rows)} 行")
    
    if not all_table_data:
        print("未找到包含labcode的表格")
        return 0
    
    print(f"\n总共收集到 {len(all_table_data)} 行数据")
    
    # 第二步：按labcode排序并合并
    # 先找到表头
    header_row = None
    data_rows = []
    
    for row in all_table_data:
        row_text = ' '.join(row).lower()
        if 'labcode' in row_text:
            if header_row is None:
                header_row = row  # 第一个表头
            else:
                # 跳过重复的表头
                continue
        else:
            data_rows.append(row)
    
    if header_row is None:
        print("未找到表头行")
        return 0
    
    # 按labcode排序数据行
    def get_labcode_value(row):
        if not row:
            return float('inf')
        labcode = row[0].strip()
        # 提取数字
        match = re.search(r'(\d+(?:\.\d+)?)', labcode)
        if match:
            try:
                return float(match.group(1))
            except:
                return float('inf')
        return float('inf')
    
    # 过滤掉labcode不是数字的行
    valid_data_rows = []
    for row in data_rows:
        labcode_val = get_labcode_value(row)
        if labcode_val != float('inf'):
            valid_data_rows.append(row)
    
    # 按labcode排序
    sorted_rows = sorted(valid_data_rows, key=get_labcode_value)
    
    # 创建完整表格
    final_table = [header_row] + sorted_rows
    
    # 第三步：保存
    df = pd.DataFrame(final_table)
    
    # 生成文件名
    header_text = ' '.join(header_row)
    if 'Eu-152' in header_text:
        table_name = "Eu152_All_Labcodes"
    elif 'Po-210' in header_text:
        table_name = "Po210_All_Labcodes"
    elif 'Pb-210' in header_text:
        table_name = "Pb210_All_Labcodes"
    elif 'Total U' in header_text or 'Total uranium' in header_text.lower():
        table_name = "TotalU_All_Labcodes"
    else:
        table_name = "All_Labcodes"
    
    # 获取labcode范围
    if sorted_rows:
        first_labcode = sorted_rows[0][0] if sorted_rows[0] else "?"
        last_labcode = sorted_rows[-1][0] if sorted_rows[-1] else "?"
        range_str = f"{extract_labcode_from_row(sorted_rows[0])}-{extract_labcode_from_row(sorted_rows[-1])}"
    else:
        range_str = "0-0"
    
    safe_name = re.sub(r'[\\/*?:"<>|]', "_", table_name)
    excel_path = os.path.join(output_folder, f"{safe_name}_{range_str}.xlsx")
    
    df.to_excel(excel_path, index=False, header=False)
    
    print(f"\n保存合并后的表格: {table_name}")
    print(f"  总行数: {len(final_table)}")
    print(f"  数据行数: {len(sorted_rows)}")
    print(f"  Labcode范围: {range_str}")
    print(f"  保存到: {excel_path}")
    
    return 1


if __name__ == "__main__":
    word_file_path = "IAEA-TERC-2021-01_summary_report.docx"
    
    if os.path.exists(word_file_path):
        print("=" * 60)
        print("开始提取并合并包含Labcode的表格")
        print("=" * 60)
        
        # 使用方法1：智能合并
        print("\n使用方法1：智能合并（按labcode顺序）")
        try:
            num_tables = extract_and_merge_labcode_tables(word_file_path)
            print(f"✓ 成功提取 {num_tables} 个表格")
        except Exception as e:
            print(f"方法1出错: {e}")
            print("\n尝试使用方法2：简单合并")
            num_tables = simple_merge_by_labcode(word_file_path)
            print(f"✓ 成功合并为 {num_tables} 个表格")
        
        print("=" * 60)
        print("处理完成！")
        
    else:
        print(f"错误: 文件 {word_file_path} 不存在")